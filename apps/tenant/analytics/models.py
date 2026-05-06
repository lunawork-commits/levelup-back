from django.core.exceptions import ValidationError
from django.core.validators import MinValueValidator
from django.db import models
from colorfield.fields import ColorField

from apps.shared.base import TimeStampedModel


# ── KnowledgeBase ─────────────────────────────────────────────────────────────

class KnowledgeBaseDocument(TimeStampedModel):
    """
    Документ базы знаний для ИИ-анализа отзывов.

    Загружается Word (.docx) или текстовый файл с инструкциями по анализу.
    Текст автоматически извлекается и используется как system-контекст
    при каждом обращении к Claude для анализа тональности отзывов.
    """

    title = models.CharField('Название', max_length=255)
    file  = models.FileField(
        'Файл (.docx / .txt)',
        upload_to='knowledge_base/',
        help_text='Загрузите Word-документ (.docx) или текстовый файл (.txt) с инструкциями.',
    )
    extracted_text = models.TextField(
        'Извлечённый текст',
        blank=True,
        help_text='Заполняется автоматически при сохранении файла.',
    )
    is_active = models.BooleanField(
        'Используется в анализе',
        default=True,
        help_text='Только активные документы передаются ИИ в качестве инструкций.',
    )

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        # Extract text after file is saved
        if self.file and not self.extracted_text:
            try:
                self.extracted_text = _extract_document_text(self.file.path)
                type(self).objects.filter(pk=self.pk).update(extracted_text=self.extracted_text)
            except Exception:
                pass

    def __str__(self):
        return self.title

    class Meta:
        verbose_name = 'Документ базы знаний'
        verbose_name_plural = 'База знаний (инструкции для ИИ)'
        ordering = ['-created_at']


def _extract_document_text(file_path: str) -> str:
    """Extract plain text from .docx or .txt file."""
    if file_path.endswith('.docx'):
        import docx
        doc = docx.Document(file_path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    with open(file_path, encoding='utf-8', errors='ignore') as f:
        return f.read()


# ── RFSegment ─────────────────────────────────────────────────────────────────

class RFSegment(TimeStampedModel):
    """
    Справочник RF-сегментов — обычно 12 штук с разными границами R и F.

    Каждый сегмент задаёт диапазоны:
      recency_min/max — дни с последнего визита
      frequency_min/max — кол-во визитов за период анализа

    Гость попадает в сегмент, если его recency_days и frequency
    оба попадают в соответствующие диапазоны.

    Записи могут быть глобальными (branch=NULL) или per-branch.
    Lookup идёт по приоритету: настройки точки → глобальные → нет данных.
    """

    branch = models.ForeignKey(
        'branch.Branch',
        on_delete=models.CASCADE,
        related_name='rf_segments',
        null=True,
        blank=True,
        verbose_name='Торговая точка',
        help_text=(
            'Оставьте пустым, чтобы создать общий сегмент для всей сети. '
            'Если у точки задан собственный сегмент с тем же кодом — используется он, '
            'иначе берётся общий.'
        ),
    )
    code = models.CharField(max_length=10, verbose_name='Код')
    name = models.CharField(max_length=100, verbose_name='Название')

    # ── Recency boundaries (days since last visit) ────────────────────────────

    recency_min = models.IntegerField(
        default=0,
        validators=[MinValueValidator(0)],
        verbose_name='Давность: от (дней)',
        help_text='Нижняя граница дней с последнего визита (включительно).',
    )
    recency_max = models.IntegerField(
        default=9999,
        validators=[MinValueValidator(0)],
        verbose_name='Давность: до (дней)',
        help_text='Верхняя граница дней с последнего визита (включительно).',
    )

    # ── Frequency boundaries (number of visits) ───────────────────────────────

    frequency_min = models.IntegerField(
        default=0,
        validators=[MinValueValidator(0)],
        verbose_name='Частота: от (визитов)',
        help_text='Минимальное кол-во визитов за период (включительно).',
    )
    frequency_max = models.IntegerField(
        default=9999,
        validators=[MinValueValidator(0)],
        verbose_name='Частота: до (визитов)',
        help_text='Максимальное кол-во визитов за период (включительно).',
    )

    # ── Display ───────────────────────────────────────────────────────────────

    emoji = models.CharField(max_length=10, verbose_name='Эмодзи')
    color = ColorField(default='#417690', verbose_name='Цвет')
    strategy = models.TextField(verbose_name='Маркетинговая стратегия')
    hint = models.TextField(
        blank=True,
        verbose_name='Подсказка для персонала',
        help_text='Краткая инструкция менеджеру. Отображается в таблице сегментов.',
    )

    # ── Campaign tracking ─────────────────────────────────────────────────────

    last_campaign_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Дата последней рассылки',
    )

    def clean(self):
        errors = {}
        if self.recency_min > self.recency_max:
            errors['recency_min'] = 'Нижняя граница давности не может быть больше верхней.'
        if self.frequency_min > self.frequency_max:
            errors['frequency_min'] = 'Нижняя граница частоты не может быть больше верхней.'
        if errors:
            raise ValidationError(errors)

    @property
    def is_global(self) -> bool:
        """True для общего сегмента (branch=None)."""
        return self.branch_id is None

    @property
    def scope_label(self) -> str:
        """Человекочитаемая область применения."""
        return 'Все точки' if self.is_global else str(self.branch)

    def __str__(self):
        return f'{self.emoji} {self.name} ({self.code}) — {self.scope_label}'

    # ── Lookup ────────────────────────────────────────────────────────────────

    @classmethod
    def resolve_for_branch_and_code(
        cls, branch_id: int | None, code: str,
    ) -> 'RFSegment | None':
        """
        Возвращает сегмент с приоритетом: настройки точки → общий.
        """
        if branch_id is not None:
            obj = cls.objects.filter(branch_id=branch_id, code=code).first()
            if obj is not None:
                return obj
        return cls.objects.filter(branch__isnull=True, code=code).first()

    # ── Mass-apply ────────────────────────────────────────────────────────────

    def _copy_defaults(self) -> dict:
        """Поля, которые копируются в per-branch версию (без branch и code)."""
        return {
            'name':           self.name,
            'recency_min':    self.recency_min,
            'recency_max':    self.recency_max,
            'frequency_min':  self.frequency_min,
            'frequency_max':  self.frequency_max,
            'emoji':          self.emoji,
            'color':          self.color,
            'strategy':       self.strategy,
            'hint':           self.hint,
        }

    def apply_to_all_branches(self) -> int:
        """
        Копирует поля сегмента (без code и branch) во все RFSegment-записи
        активных торговых точек с тем же кодом. Если у точки нет записи
        с таким кодом — создаёт.

        Возвращает количество затронутых точек.
        """
        from apps.tenant.branch.models import Branch

        defaults = self._copy_defaults()
        affected = 0
        for branch in Branch.objects.filter(is_active=True):
            type(self).objects.update_or_create(
                branch=branch, code=self.code, defaults=defaults,
            )
            affected += 1
        return affected

    def apply_to_branches(self, branch_ids: list[int]) -> int:
        """То же, что apply_to_all_branches, но только для указанных PK точек."""
        from apps.tenant.branch.models import Branch

        if not branch_ids:
            return 0

        defaults = self._copy_defaults()
        affected = 0
        for branch in Branch.objects.filter(pk__in=branch_ids, is_active=True):
            type(self).objects.update_or_create(
                branch=branch, code=self.code, defaults=defaults,
            )
            affected += 1
        return affected

    class Meta:
        verbose_name = 'RF-сегмент'
        verbose_name_plural = 'RF-сегменты'
        ordering = ['branch__name', 'recency_min', 'frequency_min']
        constraints = [
            # (branch, code) уникальны для не-NULL branch.
            models.UniqueConstraint(
                fields=['branch', 'code'],
                condition=models.Q(branch__isnull=False),
                name='analytics_rfsegment_branch_code_uniq',
            ),
            # Для глобальных сегментов (branch=NULL) уникальность по code.
            models.UniqueConstraint(
                fields=['code'],
                condition=models.Q(branch__isnull=True),
                name='analytics_rfsegment_global_code_uniq',
            ),
        ]


# ── GuestRFScore ──────────────────────────────────────────────────────────────

class GuestRFScore(models.Model):
    """
    RF-метрика гостя в конкретной торговой точке.

    Одна запись на ClientBranch — перезаписывается при каждом пересчёте.
    Пересчёт запускается через Celery (периодически или вручную).

    r_score / f_score — нормализованные баллы (1 = плохой, N = хороший).
    segment — ссылка на RFSegment, подобранный по диапазонам.
    """

    client = models.OneToOneField(
        'guest.Client',
        on_delete=models.CASCADE,
        related_name='rf_score',
        verbose_name='Гость',
    )
    recency_days = models.PositiveIntegerField(
        verbose_name='Давность (дней)',
        help_text='Дней с последнего визита на момент расчёта.',
    )
    frequency = models.PositiveIntegerField(
        verbose_name='Частота (визитов)',
        help_text='Кол-во визитов за период анализа.',
    )
    r_score = models.PositiveSmallIntegerField(
        validators=[MinValueValidator(1)],
        verbose_name='R-балл',
        help_text='Балл давности (1 = давно, выше = недавнее).',
    )
    f_score = models.PositiveSmallIntegerField(
        validators=[MinValueValidator(1)],
        verbose_name='F-балл',
        help_text='Балл частоты (1 = редко, выше = чаще).',
    )
    segment = models.ForeignKey(
        RFSegment,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='guests',
        verbose_name='Сегмент',
    )
    calculated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Рассчитано',
    )

    def __str__(self):
        seg = self.segment.code if self.segment else '—'
        return f'{self.client} [R{self.r_score} F{self.f_score} / {seg}]'

    class Meta:
        verbose_name = 'RF-метрика гостя'
        verbose_name_plural = 'RF-метрики гостей'
        indexes = [
            models.Index(fields=['r_score', 'f_score'], name='rf_score_rf_idx'),
            models.Index(fields=['segment'], name='rf_score_segment_idx'),
            models.Index(fields=['calculated_at'], name='rf_score_calc_idx'),
        ]


# ── RFMigrationLog ────────────────────────────────────────────────────────────

class RFMigrationLog(models.Model):
    """
    Журнал перемещения гостя между сегментами.

    Запись создаётся при каждом пересчёте, если сегмент изменился.
    SET_NULL на FK — чтобы не терять историю при удалении/переименовании сегмента.
    Не наследует TimeStampedModel: лог-записи никогда не обновляются,
    updated_at был бы бессмысленным полем.
    """

    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Мигрировал')

    client = models.ForeignKey(
        'guest.Client',
        on_delete=models.CASCADE,
        related_name='rf_migrations',
        verbose_name='Гость',
    )
    from_segment = models.ForeignKey(
        RFSegment,
        related_name='migrations_from',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Из сегмента',
    )
    to_segment = models.ForeignKey(
        RFSegment,
        related_name='migrations_to',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='В сегмент',
    )

    def __str__(self):
        return f'{self.client}: {self.from_segment} → {self.to_segment}'

    class Meta:
        verbose_name = 'RF-миграция'
        verbose_name_plural = 'RF-миграции'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['client', 'created_at'], name='rf_mig_client_idx'),
            models.Index(fields=['to_segment', 'created_at'], name='rf_mig_to_seg_idx'),
        ]


# ── RFSettings ────────────────────────────────────────────────────────────────

# Дефолтные пороги — единая точка истины. Используются как fallback,
# если для текущей торговой точки и для режима «Все точки» не заданы свои.
RF_DEFAULT_THRESHOLDS: dict[str, int] = {
    'r_fresh_max':    14,   # ≤ → R3 (свежий)
    'r_warm_max':     30,   # ≤ → R2 (тёплый)
    'r_cooling_max':  60,   # ≤ → R1 (остывший);  > → R0 (холодный)
    'f_rare_max':      3,   # ≤ → F1 (редко)
    'f_moderate_max':  5,   # ≤ → F2 (умеренно);  > → F3 (часто)
}


class RFSettings(TimeStampedModel):
    """
    Настройки RF-анализа.

    Одна запись на Branch ИЛИ одна запись с branch=NULL — это и есть
    «Все точки» (общие настройки сети). Пороги берутся в порядке:
      1) настройки выбранной точки (если заданы);
      2) настройки «Все точки» (branch=NULL);
      3) дефолтные значения RF_DEFAULT_THRESHOLDS.

    Пороги применяются:
      • при пересчёте RF-метрик (recalculate_rf_scores);
      • при построении RF-матрицы и отображении заголовков
        строк/колонок в клиентской админ-панели.
    """

    branch = models.OneToOneField(
        'branch.Branch',
        on_delete=models.CASCADE,
        related_name='rf_settings',
        null=True,
        blank=True,
        verbose_name='Торговая точка',
        help_text=(
            'Оставьте пустым, чтобы создать общие настройки для режима «Все точки». '
            'Допустима только одна запись с пустым полем.'
        ),
    )
    analysis_period = models.PositiveIntegerField(
        default=365,
        verbose_name='Период анализа (дней)',
        help_text='Учитываются визиты за последние N дней.',
    )
    stats_reset_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Дата обнуления статистики',
        help_text=(
            'Если задана — RF-анализ и общая статистика учитывают '
            'ТОЛЬКО данные после этой даты. '
            'Балансы монет, задания и инвентарь НЕ затрагиваются.'
        ),
    )

    # ── R-пороги (давность последнего визита, в днях) ─────────────────────────

    r_fresh_max = models.PositiveIntegerField(
        default=14,
        verbose_name='R3 «Свежий»: до (дней)',
        help_text='Гость попадает в R3, если с последнего визита прошло НЕ БОЛЕЕ этого числа дней.',
    )
    r_warm_max = models.PositiveIntegerField(
        default=30,
        verbose_name='R2 «Тёплый»: до (дней)',
        help_text='Граница R2: больше «R3 до» и не больше этого числа.',
    )
    r_cooling_max = models.PositiveIntegerField(
        default=60,
        verbose_name='R1 «Остывший»: до (дней)',
        help_text='Граница R1: больше «R2 до» и не больше этого числа. '
                  'Гости с большей давностью попадают в R0 «Холодный».',
    )

    # ── F-пороги (количество визитов за период анализа) ───────────────────────

    f_rare_max = models.PositiveIntegerField(
        default=3,
        verbose_name='F1 «Редко»: до (визитов)',
        help_text='Гость попадает в F1, если визитов НЕ БОЛЕЕ этого числа.',
    )
    f_moderate_max = models.PositiveIntegerField(
        default=5,
        verbose_name='F2 «Умеренно»: до (визитов)',
        help_text='Граница F2: больше «F1 до» и не больше этого числа. '
                  'Гости с большим числом визитов попадают в F3 «Часто».',
    )

    # ── Validation ────────────────────────────────────────────────────────────

    def clean(self):
        errors: dict[str, str] = {}

        # R-границы должны идти строго возрастающе.
        if self.r_fresh_max >= self.r_warm_max:
            errors['r_warm_max'] = 'Должно быть больше «R3 до».'
        if self.r_warm_max >= self.r_cooling_max:
            errors['r_cooling_max'] = 'Должно быть больше «R2 до».'

        # F-границы аналогично.
        if self.f_rare_max >= self.f_moderate_max:
            errors['f_moderate_max'] = 'Должно быть больше «F1 до».'

        if errors:
            raise ValidationError(errors)

    # ── Display helpers ───────────────────────────────────────────────────────

    @property
    def is_global(self) -> bool:
        """True для записи «Все точки» (branch=None)."""
        return self.branch_id is None

    @property
    def scope_label(self) -> str:
        """Человекочитаемая область применения настроек."""
        return 'Все точки' if self.is_global else str(self.branch)

    def __str__(self):
        return f'RF-настройки: {self.scope_label}'

    # ── Threshold accessors ───────────────────────────────────────────────────

    def thresholds_dict(self) -> dict[str, int]:
        """Текущие пороги в виде словаря (тот же набор ключей, что в RF_DEFAULT_THRESHOLDS)."""
        return {
            'r_fresh_max':    self.r_fresh_max,
            'r_warm_max':     self.r_warm_max,
            'r_cooling_max':  self.r_cooling_max,
            'f_rare_max':     self.f_rare_max,
            'f_moderate_max': self.f_moderate_max,
        }

    @classmethod
    def get_global(cls) -> 'RFSettings | None':
        """Запись «Все точки» (branch=NULL) или None, если её ещё не создали."""
        return cls.objects.filter(branch__isnull=True).first()

    @classmethod
    def resolve_for_scope(
        cls, branch_ids: list[int] | None,
    ) -> tuple['RFSettings | None', dict[str, int], str]:
        """
        Возвращает (settings_obj, thresholds_dict, source).

        source ∈ {'branch', 'global', 'default'} — откуда взяты пороги.

        Логика:
          • ровно одна точка в выборке и для неё есть запись → её пороги;
          • иначе берём «Все точки» (branch=NULL);
          • иначе — RF_DEFAULT_THRESHOLDS.
        """
        # 1) Точечные настройки имеют смысл только при выборе одной точки.
        if branch_ids and len(branch_ids) == 1:
            obj = cls.objects.filter(branch_id=branch_ids[0]).first()
            if obj is not None:
                return obj, obj.thresholds_dict(), 'branch'

        # 2) Общие настройки «Все точки».
        glob = cls.get_global()
        if glob is not None:
            return glob, glob.thresholds_dict(), 'global'

        # 3) Захардкоженные дефолты.
        return None, dict(RF_DEFAULT_THRESHOLDS), 'default'

    @classmethod
    def thresholds_for_scope(cls, branch_ids: list[int] | None) -> dict[str, int]:
        """Удобная обёртка: только пороги, без объекта и метаданных."""
        _, thresholds, _ = cls.resolve_for_scope(branch_ids)
        return thresholds

    # ── Mass-apply (Task 3) ───────────────────────────────────────────────────

    def apply_thresholds_to_all_branches(self) -> int:
        """
        Копирует пороги (R/F и analysis_period) в RFSettings всех активных торговых точек.

        - Для точек, у которых ещё нет RFSettings — создаёт запись.
        - Для существующих — перезаписывает значения порогов.
        - Не трогает stats_reset_date (это индивидуальная настройка).
        - Не трогает запись «Все точки» (branch=NULL): источник остаётся отдельным.

        Возвращает количество затронутых точек.
        """
        from apps.tenant.branch.models import Branch

        affected = 0
        defaults = {
            'analysis_period': self.analysis_period,
            'r_fresh_max':     self.r_fresh_max,
            'r_warm_max':      self.r_warm_max,
            'r_cooling_max':   self.r_cooling_max,
            'f_rare_max':      self.f_rare_max,
            'f_moderate_max':  self.f_moderate_max,
        }

        for branch in Branch.objects.filter(is_active=True):
            type(self).objects.update_or_create(branch=branch, defaults=defaults)
            affected += 1
        return affected

    def apply_thresholds_to_branches(self, branch_ids: list[int]) -> int:
        """То же, что и apply_thresholds_to_all_branches, но только для указанных PK."""
        from apps.tenant.branch.models import Branch

        if not branch_ids:
            return 0

        defaults = {
            'analysis_period': self.analysis_period,
            'r_fresh_max':     self.r_fresh_max,
            'r_warm_max':      self.r_warm_max,
            'r_cooling_max':   self.r_cooling_max,
            'f_rare_max':      self.f_rare_max,
            'f_moderate_max':  self.f_moderate_max,
        }
        affected = 0
        for branch in Branch.objects.filter(pk__in=branch_ids, is_active=True):
            type(self).objects.update_or_create(branch=branch, defaults=defaults)
            affected += 1
        return affected

    class Meta:
        verbose_name = 'RF-настройки'
        verbose_name_plural = 'RF-настройки'
        # Гарантируем не более одной записи «Все точки» (branch=NULL).
        # OneToOne уже даёт уникальность для не-NULL значений; добавляем
        # отдельный частичный уникальный индекс для NULL.
        constraints = [
            models.UniqueConstraint(
                fields=['branch'],
                condition=models.Q(branch__isnull=True),
                name='analytics_rfsettings_one_global',
            ),
        ]


# ── BranchSegmentSnapshot ─────────────────────────────────────────────────────

class BranchSegmentSnapshot(TimeStampedModel):
    """
    Ежедневный снапшот: сколько гостей в каждом сегменте по каждой точке.

    Хранит историю для построения трендов без пересчёта каждый раз.
    Обновляется через:
        BranchSegmentSnapshot.objects.update_or_create(
            branch=branch, segment=segment, date=today,
            defaults={'guests_count': count}
        )

    date — явно проставляется кодом (не auto_now_add) для возможности
    ретроспективного заполнения.
    """

    branch = models.ForeignKey(
        'branch.Branch',
        on_delete=models.CASCADE,
        related_name='segment_snapshots',
        verbose_name='Торговая точка',
    )
    segment = models.ForeignKey(
        RFSegment,
        on_delete=models.CASCADE,
        related_name='snapshots',
        verbose_name='Сегмент',
    )
    guests_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Кол-во гостей',
    )
    date = models.DateField(
        db_index=True,
        verbose_name='Дата',
        help_text='Дата расчёта снапшота.',
    )

    def __str__(self):
        return f'{self.date} | {self.branch.name} | {self.segment.code}: {self.guests_count}'

    class Meta:
        unique_together = ('branch', 'segment', 'date')
        ordering = ['-date', 'branch', 'segment']
        verbose_name = 'Снапшот сегмента'
        verbose_name_plural = 'Снапшоты сегментов'
        indexes = [
            models.Index(fields=['branch', 'date'], name='snapshot_branch_date_idx'),
            models.Index(fields=['segment', 'date'], name='snapshot_segment_date_idx'),
        ]


# ── Delivery RF (separate metrics, same RFSegment definitions) ─────────────────

class GuestRFScoreDelivery(models.Model):
    """
    RF-метрика гостя по активациям доставки.

    Аналог GuestRFScore, но recency/frequency считается по Delivery.activated_by,
    а не по ClientBranchVisit. Использует те же RFSegment, что и ресторанный RF.
    """

    client = models.OneToOneField(
        'guest.Client',
        on_delete=models.CASCADE,
        related_name='rf_score_delivery',
        verbose_name='Гость',
    )
    recency_days = models.PositiveIntegerField(
        verbose_name='Давность (дней)',
        help_text='Дней с последней активации доставки.',
    )
    frequency = models.PositiveIntegerField(
        verbose_name='Частота (заказов)',
        help_text='Кол-во активаций доставки за период анализа.',
    )
    r_score = models.PositiveSmallIntegerField(
        validators=[MinValueValidator(1)],
        verbose_name='R-балл',
    )
    f_score = models.PositiveSmallIntegerField(
        validators=[MinValueValidator(1)],
        verbose_name='F-балл',
    )
    segment = models.ForeignKey(
        RFSegment,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='delivery_guests',
        verbose_name='Сегмент',
    )
    calculated_at = models.DateTimeField(auto_now=True, verbose_name='Рассчитано')

    def __str__(self):
        seg = self.segment.code if self.segment else '—'
        return f'[Доставка] {self.client} [R{self.r_score} F{self.f_score} / {seg}]'

    class Meta:
        verbose_name = 'RF-метрика гостя (доставка)'
        verbose_name_plural = 'RF-метрики гостей (доставка)'
        indexes = [
            models.Index(fields=['r_score', 'f_score'], name='rf_del_score_rf_idx'),
            models.Index(fields=['segment'],             name='rf_del_segment_idx'),
            models.Index(fields=['calculated_at'],       name='rf_del_calc_idx'),
        ]


class RFMigrationLogDelivery(models.Model):
    """
    Журнал смены RF-сегментов по доставке.
    Аналог RFMigrationLog для доставочного RF.
    """

    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Мигрировал')

    client = models.ForeignKey(
        'guest.Client',
        on_delete=models.CASCADE,
        related_name='rf_migrations_delivery',
        verbose_name='Гость',
    )
    from_segment = models.ForeignKey(
        RFSegment,
        related_name='delivery_migrations_from',
        on_delete=models.SET_NULL,
        null=True, blank=True,
        verbose_name='Из сегмента',
    )
    to_segment = models.ForeignKey(
        RFSegment,
        related_name='delivery_migrations_to',
        on_delete=models.SET_NULL,
        null=True, blank=True,
        verbose_name='В сегмент',
    )

    def __str__(self):
        return f'[Доставка] {self.client}: {self.from_segment} → {self.to_segment}'

    class Meta:
        verbose_name = 'RF-миграция (доставка)'
        verbose_name_plural = 'RF-миграции (доставка)'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['client', 'created_at'],     name='rf_del_mig_client_idx'),
            models.Index(fields=['to_segment', 'created_at'], name='rf_del_mig_to_seg_idx'),
        ]


class POSGuestCache(models.Model):
    """
    Кэш количества гостей из POS-системы (IIKO / Dooglys).

    Одна запись на (branch, date) — перезаписывается Celery-задачей
    fetch_pos_data_all_tenants_task, которая запускается ежедневно.

    get_pos_guests_count() суммирует эти записи за нужный диапазон дат
    вместо прямого обращения к POS API при каждом запросе дашборда.
    """

    branch = models.ForeignKey(
        'branch.Branch',
        on_delete=models.CASCADE,
        related_name='pos_guest_cache',
        verbose_name='Торговая точка',
    )
    date = models.DateField(db_index=True, verbose_name='Дата')
    guest_count = models.PositiveIntegerField(default=0, verbose_name='Кол-во гостей')
    fetched_at = models.DateTimeField(auto_now=True, verbose_name='Обновлено')

    class Meta:
        unique_together = ('branch', 'date')
        verbose_name = 'Кэш POS гостей'
        verbose_name_plural = 'Кэш POS гостей'
        indexes = [
            models.Index(fields=['branch', 'date'], name='pos_cache_branch_date_idx'),
        ]

    def __str__(self):
        return f'{self.date} | {self.branch.name}: {self.guest_count}'


class BranchSegmentSnapshotDelivery(TimeStampedModel):
    """
    Ежедневный снапшот распределения гостей по сегментам доставки.
    Аналог BranchSegmentSnapshot для доставочного RF.
    """

    branch = models.ForeignKey(
        'branch.Branch',
        on_delete=models.CASCADE,
        related_name='segment_snapshots_delivery',
        verbose_name='Торговая точка',
    )
    segment = models.ForeignKey(
        RFSegment,
        on_delete=models.CASCADE,
        related_name='delivery_snapshots',
        verbose_name='Сегмент',
    )
    guests_count = models.PositiveIntegerField(default=0, verbose_name='Кол-во гостей')
    date = models.DateField(db_index=True, verbose_name='Дата')

    def __str__(self):
        return f'[Дост.] {self.date} | {self.branch.name} | {self.segment.code}: {self.guests_count}'

    class Meta:
        unique_together = ('branch', 'segment', 'date')
        ordering = ['-date', 'branch', 'segment']
        verbose_name = 'Снапшот сегмента (доставка)'
        verbose_name_plural = 'Снапшоты сегментов (доставка)'
        indexes = [
            models.Index(fields=['branch', 'date'],   name='snap_del_branch_date_idx'),
            models.Index(fields=['segment', 'date'],  name='snap_del_segment_date_idx'),
        ]
